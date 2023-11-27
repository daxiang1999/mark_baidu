from docx import Document
import pandas as pd
import numpy as np


def get_docx_table(docx_path: str, table_index: int = 0) -> [dict]:
    """获取docx文件中的表格"""
    document = Document(docx_path)
    table = document.tables[table_index]
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return data


def get_docx_table_df(docx_path: str, table_index: int = 0) -> pd.DataFrame:
    """获取docx文件中的表格，并转换成DataFrame"""
    data = get_docx_table(docx_path, table_index)
    return pd.DataFrame(data)


def merge_word_files(doc1: Document, doc2: Document) -> Document:
    """
    合并两个word文档，将第二个文件的内容追加到第一个文件的末尾
    :return:
    """
    for element in doc2.element.body:
        doc1.element.body.append(element)

    return doc1


def delete_after_paragraph(document: Document, paragraph_text: str) -> Document:
    """
    删除指定段落之后的所有段落和表格内容
    :param document: docx.document.Document 对象
    :param paragraph_text: 指定的段落文本
    :return:
    """
    for i, element in enumerate(document.element.body):
        tag = element.tag.split('}')[1]
        if tag == 'p':  # 如果是段落
            paragraph = document.element.body[i]
            if paragraph.text == paragraph_text:
                total_length = len(document.element.body)
                # 这里remove_index是指定段落之后的第一个元素的索引。如果删除指定段落之后的所有内容（包括指定段落），remove_index=i
                remove_index = i + 1
                for _ in range(remove_index, total_length):
                    document.element.body[remove_index].getparent().remove(document.element.body[remove_index])
                break

    return document


# 在指定内容之后插入一个段落//TODO: 未测试还有bug
def insert_paragraph(document: Document, paragraph_text: str, insert_text: str):
    """
    在指定内容之后插入一个段落
    :param document: docx.document.Document 对象
    :param paragraph_text: 指定的段落文本
    :param insert_text: 要插入的段落文本
    :return:
    """
    for i, element in enumerate(document.element.body):
        tag = element.tag.split('}')[1]
        if tag == 'p':  # 如果是段落
            paragraph = document.element.body[i]
            if paragraph.text == paragraph_text:
                paragraph.insert_paragraph_before(insert_text)
                break

    return document


# 将DataFrame数据写入word指定表格  //TODO: 数据中还有
def write_df_to_word(df, doc):
    """
    将DataFrame数据写入word指定表格
    :param df: DataFrame数据
    :param doc: docx.document.Document 对象
    :return:
    """
    # 创建表格
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')
    # 写入表头
    for j, value in enumerate(df.columns):
        table.cell(0, j).text = str(value)
    # 写入数据
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.iloc[i, j])
    return doc


if __name__ == '__main__':
    path = r"C:\Users\xuxx0\Desktop\test.docx"
    doc = Document(path)
    df = pd.DataFrame(np.arange(16).reshape(4, 4), columns=list('abcd'))
    doc = write_df_to_word(df, doc)
    doc.save(path.replace('test', 'test2'))

